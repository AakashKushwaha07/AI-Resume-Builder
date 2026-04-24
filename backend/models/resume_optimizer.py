import uuid
import json
import html as html_lib
import re
import os
from datetime import datetime, timedelta
from io import BytesIO
from typing import Dict, List, Any, Optional
from urllib.parse import quote_plus
from docx import Document
from docx.shared import Pt
import fitz
import requests
from models.semantic_matcher import semantic_match, semantic_text_similarity
from models.evaluation_engine import evaluate_resume_against_jd
from utils.resume_json_parser import parse_resume_to_json
from utils.jd_json_parser import parse_jd_to_json

class ResumeOptimizerAgent:
    """
    AI Agent for conversational resume optimization
    Provides personalized recommendations and helps users improve their resumes
    """

    def __init__(self):
        self.sessions = {}  # session_id -> session_data
        self.session_timeout = timedelta(hours=2)  # Sessions expire after 2 hours

        # Predefined optimization templates and suggestions
        self.templates = {
            'summary': {
                'structure': 'Start with your current role, years of experience, key skills, and career goals.',
                'keywords': ['experienced', 'skilled', 'results-driven', 'proven track record'],
                'length': '2-3 sentences, 50-100 words'
            },
            'experience': {
                'structure': 'Use action verbs, quantify achievements, focus on relevant experience.',
                'action_verbs': ['developed', 'implemented', 'managed', 'created', 'optimized', 'led'],
                'quantifiers': ['increased by', 'reduced by', 'managed team of', 'processed']
            },
            'skills': {
                'categorization': ['Technical Skills', 'Soft Skills', 'Tools & Technologies'],
                'relevance': 'Prioritize skills mentioned in job description'
            }
        }

        self.intent_examples = {
            'weakness_analysis': [
                'what are the weaknesses in this resume',
                'what is missing from this resume',
                'where is this resume weak',
            ],
            'score_explanation': [
                'how is the resume score generated',
                'what is the score based on',
                'why is the score this value',
            ],
            'resume_rewrite': [
                'rewrite my resume for this job description',
                'tailor my resume to the jd',
                'optimize my resume for this role',
            ],
            'web_search': [
                'search the web for the latest information',
                'what is the current market trend',
                'find recent news about this topic',
            ],
            'summary': [
                'improve my professional summary',
                'rewrite my summary',
                'objective section help',
            ],
            'career_readiness': [
                'am i job ready',
                'what should i learn next',
                'how can i become more job ready',
                'is my resume good',
                'how can i improve my resume',
            ],
        }
        self.intent_threshold = 0.42
        self.resume_generation_policy = (
            "Generate only the final resume content. Preserve original resume facts and structure, "
            "avoid placeholders, avoid internal notes, avoid generic filler, and never fabricate experience."
        )

    def start_session(self, resume_json: Dict[str, Any], jd_json: Optional[Dict[str, Any]] = None) -> str:
        """Start a new optimization session"""
        session_id = str(uuid.uuid4())

        session_data = {
            'session_id': session_id,
            'resume_json': resume_json,
            'jd_json': jd_json,
            'conversation_history': [],
            'suggestions_given': [],
            'sections_optimized': [],
            'created_at': datetime.now(),
            'last_activity': datetime.now()
        }

        self.sessions[session_id] = session_data
        return session_id

    def analyze_resume(self, resume_json: Dict[str, Any], jd_json: Optional[Dict[str, Any]] = None) -> str:
        """Analyze resume and provide initial feedback"""
        try:
            analysis = "📊 **Initial Resume Analysis:**\n\n"

            # Basic stats
            experience = resume_json.get('experience_years', 0)
            skills = resume_json.get('tech_skills', [])
            education = resume_json.get('education', [])

            analysis += f"**Resume Summary:**\n"
            analysis += f"• Experience: {experience} years\n"
            analysis += f"• Technical Skills: {len(skills)} listed\n"
            analysis += f"• Education: {', '.join(education) if education else 'Not specified'}\n\n"

            if jd_json:
                # Job-specific analysis
                from models.evaluation_engine import evaluate_resume_against_jd
                evaluation = evaluate_resume_against_jd(resume_json, jd_json)

                analysis += f"**Job Match Score: {evaluation['final_score']:.1f}/100**\n"
                analysis += f"Status: {evaluation['eligibility']}\n\n"

                analysis += "**Matching Criteria:**\n"
                skills_detail = evaluation['details']['skills']['mandatory']
                analysis += f"• Skills: {len(skills_detail['matched'])}/{len(skills_detail['matched']) + len(skills_detail['missing'])} matched\n"

                exp_detail = evaluation['details']['experience']
                analysis += f"• Experience: {exp_detail['found_years']} years (required: {exp_detail['required_years']})\n"

                if not evaluation['eligibility'] == 'PASS':
                    analysis += "\n💡 **To improve your match:**\n"
                    if len(skills_detail['missing']) > 0:
                        missing = list(skills_detail['missing'])[:3]
                        analysis += f"• Learn these skills: {', '.join(missing)}\n"
                    if exp_detail['gap_years'] > 0:
                        analysis += f"• Gain {exp_detail['gap_years']} more years of experience\n"

            return analysis

        except Exception as e:
            return f"Quick Resume Check:\n• Experience: {resume_json.get('experience_years', 0)} years\n• Skills: {len(resume_json.get('tech_skills', []))} identified\n\nReady to optimize! What aspect would you like to improve?"

    def process_message(self, session_id: str, user_message: str) -> str:
        """Process user message and return AI response"""
        if session_id not in self.sessions:
            return "Session expired. Please start a new optimization session."

        session = self.sessions[session_id]
        session['last_activity'] = datetime.now()

        # Add user message to conversation history
        session['conversation_history'].append({
            'role': 'user',
            'message': user_message,
            'timestamp': datetime.now()
        })

        # Analyze the message and generate response
        response = self._generate_response(session, user_message.lower())

        # Add AI response to conversation history
        session['conversation_history'].append({
            'role': 'assistant',
            'message': response,
            'timestamp': datetime.now()
        })

        return response

    def _format_bullets(self, items: List[str]) -> str:
        return "\n".join(f"- {item}" for item in items if item)

    def _get_evaluation(self, session: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        resume = session['resume_json']
        jd = session.get('jd_json')
        if not jd:
            return None
        try:
            return evaluate_resume_against_jd(resume, jd)
        except Exception:
            return None

    def _load_semantic_state(self) -> Dict[str, Any]:
        """Templates have been removed; keep a neutral empty state."""
        return {"active_template_id": None, "selected_template": None}

    def _template_profile(self, template_id: str) -> Dict[str, Any]:
        """Return a neutral semantic-only writing profile."""
        return {
            "name": "Semantic Assistant",
            "section_order": ["summary", "experience", "skills", "projects", "education"],
            "summary_style": "concise",
            "skills_style": "balanced",
            "tone": "semantic and job-aware",
        }

    def _looks_like_full_resume_rewrite(self, message: str) -> bool:
        rewrite_words = ['edit my resume', 'rewrite my resume', 'update my resume', 'optimize my resume', 'tailor my resume']
        jd_words = ['job description', 'jd', 'job role', 'match', 'matching criteria']
        return any(word in message for word in rewrite_words) and any(word in message for word in jd_words)

    def _wants_weakness_analysis(self, message: str) -> bool:
        keywords = [
            'weakness', 'weaknesses', 'gap', 'gaps', 'missing', 'what is wrong',
            'what is missing', 'resume issues', 'improve this resume'
        ]
        return any(word in message for word in keywords)

    def _wants_score_explanation(self, message: str) -> bool:
        keywords = [
            'score generated', 'score is generated', 'how is the score',
            'basis of score', 'on what bases', 'what is score based on',
            'why this score', 'resume score'
        ]
        return any(word in message for word in keywords)

    def _semantic_intent_match(self, message: str) -> tuple[str, float]:
        best_intent = 'general'
        best_score = 0.0

        for intent, examples in self.intent_examples.items():
            for example in examples:
                score = semantic_text_similarity(message, example)
                if score > best_score:
                    best_intent = intent
                    best_score = score

        return best_intent, best_score

    def _looks_like_web_question(self, message: str) -> bool:
        web_markers = [
            'latest', 'current', 'today', 'news', 'recent', 'this year',
            'web search', 'search the web', 'look up', 'find online', 'official source'
        ]
        return any(marker in message for marker in web_markers)

    def _search_web(self, query: str, max_results: int = 3) -> List[Dict[str, str]]:
        search_url = f"https://html.duckduckgo.com/html/?q={quote_plus(query)}"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36'
        }

        response = requests.get(search_url, headers=headers, timeout=10)
        response.raise_for_status()
        html_text = response.text

        results: List[Dict[str, str]] = []
        pattern = re.compile(
            r'<a[^>]*class="result__a"[^>]*href="(?P<href>[^"]+)"[^>]*>(?P<title>.*?)</a>.*?'
            r'<a[^>]*class="result__snippet"[^>]*>(?P<snippet>.*?)</a>',
            re.S,
        )
        for match in pattern.finditer(html_text):
            title = re.sub(r'<.*?>', '', html_lib.unescape(match.group('title'))).strip()
            snippet = re.sub(r'<.*?>', '', html_lib.unescape(match.group('snippet'))).strip()
            href = html_lib.unescape(match.group('href')).strip()
            results.append({
                'title': title,
                'snippet': snippet,
                'url': href,
            })
            if len(results) >= max_results:
                break

        return results

    def _build_targeted_rewrite_response(self, session: Dict[str, Any], message: str) -> str:
        evaluation = self._get_evaluation(session)
        if not evaluation:
            return (
                "I can tailor your resume, but I need the job description in the session to do that well.\n\n"
                "Once you share the JD, I will:\n"
                "- identify the strongest matching skills\n"
                "- call out the real gaps\n"
                "- suggest exact edits for summary, skills, experience, and projects"
            )

        skills_detail = evaluation['details']['skills']['mandatory']
        exp_detail = evaluation['details']['experience']
        matched = skills_detail['matched']
        missing = skills_detail['missing']

        response_parts = [
            "I can tailor your resume to this job description and make it stronger.",
            "I will keep this grounded in your actual resume and avoid fabricating experience or skills.",
            "",
            f"Right now your estimated match is {evaluation['final_score']:.1f}/100.",
        ]

        strengths = []
        if matched:
            strengths.append(f"Already aligned skills: {', '.join(matched[:5])}")
        if exp_detail['found_years'] > 0:
            strengths.append(f"Experience detected: {exp_detail['found_years']} years")
        if session['resume_json'].get('projects'):
            strengths.append("Projects section is available and can be used to strengthen relevance")

        gaps = []
        if missing:
            gaps.append(f"Missing skills to address carefully: {', '.join(missing[:5])}")
        if exp_detail['gap_years'] > 0:
            gaps.append(
                f"Experience gap: job asks for {exp_detail['required_years']} years and the resume currently shows {exp_detail['found_years']}"
            )

        if strengths:
            response_parts.extend([
                "",
                "The strongest parts to lean into are:",
                self._format_bullets(strengths),
            ])

        if gaps:
            response_parts.extend([
                "",
                "The main gaps to address carefully are:",
                self._format_bullets(gaps),
            ])

        next_steps = [
            "rewrite the summary around the target role",
            "bring the most relevant skills and projects higher",
            "tighten the experience bullets around JD keywords",
            "strengthen project bullets with clear impact",
        ]
        response_parts.extend([
            "",
            "My next move would be:",
            self._format_bullets(next_steps),
        ])

        return "\n".join(response_parts)

    def _build_career_readiness_response(self, session: Dict[str, Any]) -> str:
        resume = session['resume_json']
        jd = session['jd_json']
        evaluation = self._get_evaluation(session)

        response = ["Here is your career-readiness snapshot:"]

        if evaluation:
            response.append(f"- ATS-style match: {evaluation['final_score']:.1f}/100")
            response.append(f"- Eligibility signal: {evaluation['eligibility']}")
        else:
            response.append(f"- Experience detected: {resume.get('experience_years', 0)} years")
            response.append(f"- Skills detected: {len(resume.get('tech_skills', []))}")

        missing_skills = []
        if jd:
            jd_skills = jd.get('mandatory_skills', []) + jd.get('preferred_skills', [])
            resume_skills = set(resume.get('tech_skills', []))
            missing_skills = [skill for skill in jd_skills if skill not in resume_skills]

        response.append("")
        response.append("What to work on next:")

        if missing_skills:
            response.append(f"- Skills: learn or demonstrate {', '.join(missing_skills[:5])}")
        else:
            response.append("- Skills: current alignment looks reasonable; focus on depth and proof")

        if not resume.get('summary'):
            response.append("- Summary: add a short role-focused summary at the top")
        else:
            response.append("- Summary: make it more role-specific and outcome-driven")

        if not resume.get('projects'):
            response.append("- Projects: add 1 to 2 proof-of-work projects if possible")
        else:
            response.append("- Projects: rewrite bullets with impact, tools, and outcomes")

        return "\n".join(response)

    def _is_job_fit_question(self, message: str) -> bool:
        lowercase = message.lower()
        return any(
            phrase in lowercase for phrase in [
                'is my resume good',
                'good for this role',
                'fit for this role',
                'job ready',
                'ready for this role',
                'is this resume good',
                'suitable for this role'
            ]
        )

    def _is_improvement_question(self, message: str) -> bool:
        lowercase = message.lower()
        return any(
            phrase in lowercase for phrase in [
                'how can i improve',
                'how to improve',
                'what can i improve',
                'improve my resume',
                'make my resume better',
                'what should i learn next',
                'how can i improve my resume'
            ]
        )

    def _build_resume_fit_response(self, session: Dict[str, Any]) -> str:
        evaluation = self._get_evaluation(session)
        if not evaluation:
            return (
                "I need the job description to tell you whether this resume is a fit for the role. "
                "Please provide the target job or role so I can compare it directly."
            )

        response = [
            "Resume fit check:",
            f"- Current score: {evaluation['final_score']:.1f}/100",
            f"- Eligibility: {evaluation['eligibility']}"
        ]

        issues = []
        skills_detail = evaluation['details']['skills']['mandatory']
        missing_skills = skills_detail.get('missing', [])
        if missing_skills:
            issues.append(f"Missing mandatory skills: {', '.join(missing_skills[:5])}")

        exp_detail = evaluation['details']['experience']
        if exp_detail.get('gap_years', 0) > 0:
            issues.append(
                f"Experience gap: {exp_detail['gap_years']} years below the job requirement of {exp_detail['required_years']} years"
            )

        projects = evaluation['details']['projects']
        if projects.get('required') and not projects.get('pass'):
            issues.append("Add or highlight at least one relevant project with measurable outcomes")

        education = evaluation['details']['education']
        if not education.get('pass'):
            issues.append("Align your education section to the job requirements")

        if not issues:
            response.append("")
            response.append("Your resume is a solid fit for this role. Focus on stronger language and proof points in the experience bullets.")
        else:
            response.append("")
            response.append("The top improvements to make are:")
            response.extend(self._format_bullets(issues[:4]))
            response.append("")
            response.append("Once you fix these gaps, I can help rewrite the summary, experience bullets, or skills section.")

        return "\n".join(response)

    def _build_improvement_response(self, session: Dict[str, Any]) -> str:
        resume = session['resume_json']
        evaluation = self._get_evaluation(session)
        response = ["How to improve your resume:"]

        if evaluation:
            skills_detail = evaluation['details']['skills']['mandatory']
            missing_skills = skills_detail.get('missing', [])
            exp_detail = evaluation['details']['experience']
            projects = evaluation['details']['projects']

            if missing_skills:
                response.append(f"- Add or better demonstrate these required skills: {', '.join(missing_skills[:5])}")
            if exp_detail.get('gap_years', 0) > 0:
                response.append(
                    f"- Close the experience gap by emphasizing relevant projects or applicable coursework for the required {exp_detail['required_years']} years"
                )
            if projects.get('required') and not projects.get('pass'):
                response.append("- Add one or two strong project bullets that show results and tools used")
        else:
            response.append("- I can give more targeted advice once you share the target job description.")

        if not resume.get('summary'):
            response.append("- Add a concise, role-focused professional summary at the top")
        else:
            response.append("- Make your summary more specific to the target role and outcomes")

        if resume.get('tech_skills'):
            response.append("- Reorder your skills so the most relevant technologies appear first")
        else:
            response.append("- Add a dedicated skills section with the most job-relevant technologies")

        if not resume.get('projects'):
            response.append("- Include at least one proof-of-work project or achievement bullet")

        response.append("")
        response.append("Next, pick one section and I’ll help you rewrite it: summary, experience, skills, or projects.")

        return "\n".join(response)

    def _build_weakness_response(self, session: Dict[str, Any]) -> str:
        resume = session['resume_json']
        evaluation = self._get_evaluation(session)

        weaknesses = []
        if not resume.get('summary'):
            weaknesses.append("No professional summary is present, so the resume loses impact at the top.")

        if evaluation:
            exp_detail = evaluation['details']['experience']
            skills_detail = evaluation['details']['skills']['mandatory']
            edu_detail = evaluation['details']['education']
            project_detail = evaluation['details']['projects']

            if exp_detail['gap_years'] > 0:
                weaknesses.append(
                    f"Experience is below the JD requirement by about {exp_detail['gap_years']} years."
                )
            if skills_detail['missing']:
                weaknesses.append(
                    f"Some important JD skills are not clearly shown: {', '.join(skills_detail['missing'][:5])}."
                )
            if project_detail['required'] and not project_detail['found']:
                weaknesses.append("The job expects projects, but the resume does not currently show them clearly.")
            if not edu_detail['pass']:
                weaknesses.append("Education criteria do not fully align with the job description.")
        else:
            if not resume.get('tech_skills'):
                weaknesses.append("Technical skills are not clearly highlighted.")
            if not resume.get('projects'):
                weaknesses.append("Projects are missing or not easy to detect.")
            if not resume.get('experience_years'):
                weaknesses.append("Experience is not clearly quantified.")

        if not weaknesses:
            weaknesses.append("The resume is fairly balanced, but it would still benefit from stronger wording and more measurable achievements.")

        return (
            "Here are the main weaknesses to focus on first:\n\n"
            f"{self._format_bullets(weaknesses)}\n\n"
            "If you want, I can turn this into a section-by-section rewrite plan."
        )

    def _build_score_explanation_response(self, session: Dict[str, Any]) -> str:
        resume = session['resume_json']
        evaluation = self._get_evaluation(session)

        if evaluation:
            skills = evaluation['details']['skills']
            education = evaluation['details']['education']
            experience = evaluation['details']['experience']
            projects = evaluation['details']['projects']

            return (
                "The score is based on how well the resume matches the job and how complete it looks overall:\n\n"
                f"- Education fit: {'pass' if education['pass'] else 'fail'}\n"
                f"- Experience fit: {'pass' if experience['pass'] else 'fail'}\n"
                f"- Mandatory skill match: {skills['mandatory']['match_ratio']:.2f}%\n"
                f"- Project requirement: {'pass' if projects['pass'] else 'fail'}\n"
                f"- Preferred and soft skills add bonus points\n\n"
                f"Your current score is {evaluation['final_score']:.1f}/100."
            )

        return (
            "The resume score is a blended quality score based on:\n\n"
            "- education\n"
            "- experience\n"
            "- technical skills\n"
            "- soft skills\n"
            "- projects or proof of work\n\n"
            "When a job description is available, I score it more specifically against the JD. "
            f"Right now I can see {len(resume.get('tech_skills', []))} technical skills, "
            f"{len(resume.get('education', []))} education entries, and "
            f"{len(resume.get('projects', []))} project entries."
        )

    def _build_web_search_response(self, message: str) -> str:
        try:
            results = self._search_web(message)
        except Exception as e:
            return (
                "I tried to look that up online, but the search request didn’t come back cleanly.\n\n"
                f"Error: {str(e)}\n\n"
                "If you want, I can still answer from the resume/JD context or try a different query."
            )

        if not results:
            return (
                "I searched the web, but I couldn’t find a strong result for that question.\n\n"
                "Try rephrasing it a little more specifically, and I’ll look again."
            )

        lines = ["Here’s what I found online:"]
        for item in results:
            lines.append(f"- {item['title']}: {item['snippet']}")
        lines.append("")
        lines.append("If you want, I can also turn this into a resume/JD-specific recommendation.")
        return "\n".join(lines)

    def _generate_response(self, session: Dict[str, Any], message: str) -> str:
        """Generate contextual response based on user message"""

        intent, semantic_score = self._semantic_intent_match(message)

        if self._is_job_fit_question(message):
            return self._build_resume_fit_response(session)

        if self._is_improvement_question(message):
            return self._build_improvement_response(session)

        if self._looks_like_web_question(message) or (intent == 'web_search' and semantic_score >= self.intent_threshold):
            return self._build_web_search_response(message)

        if self._looks_like_full_resume_rewrite(message) or (intent == 'resume_rewrite' and semantic_score >= self.intent_threshold):
            return self._build_targeted_rewrite_response(session, message)

        if self._wants_weakness_analysis(message) or (intent == 'weakness_analysis' and semantic_score >= self.intent_threshold):
            return self._build_weakness_response(session)

        if self._wants_score_explanation(message) or (intent == 'score_explanation' and semantic_score >= self.intent_threshold):
            return self._build_score_explanation_response(session)

        if intent == 'career_readiness' and semantic_score >= self.intent_threshold:
            return self._build_career_readiness_response(session)

        # Keywords for different types of requests
        if any(word in message for word in ['summary', 'objective', 'profile']):
            return self._handle_summary_request(session)

        elif any(word in message for word in ['experience', 'work', 'job']):
            return self._handle_experience_request(session)

        elif any(word in message for word in ['skills', 'technologies', 'competencies']):
            return self._handle_skills_request(session)

        elif any(word in message for word in ['education', 'degree', 'qualification']):
            return self._handle_education_request(session)

        elif any(word in message for word in ['keyword', 'ats', 'optimization']):
            return self._handle_keyword_optimization(session)

        elif any(word in message for word in ['rewrite', 'improve', 'better']):
            return self._handle_rewrite_request(session, message)

        elif any(word in message for word in ['analyze', 'review', 'feedback']):
            return self._handle_analysis_request(session)

        else:
            return self._handle_general_request(session, message)

    def _handle_summary_request(self, session: Dict[str, Any]) -> str:
        """Handle requests about resume summary/objective"""
        resume = session['resume_json']
        jd = session['jd_json']

        response = "📝 **Resume Summary Optimization:**\n\n"

        # Check if summary exists
        current_summary = resume.get('summary', '')
        if not current_summary:
            response += "You don't have a professional summary. This is crucial for ATS systems!\n\n"
        else:
            response += f"Your current summary: \"{current_summary[:100]}...\"\n\n"

        # Provide template and suggestions
        response += "**Recommended Structure:**\n"
        response += "• Start with your current role and years of experience\n"
        response += "• Highlight 2-3 key skills relevant to the target job\n"
        response += "• Include quantifiable achievements\n"
        response += "• End with your career goal\n\n"

        if jd:
            # Suggest keywords from job description
            jd_skills = jd.get('mandatory_skills', []) + jd.get('preferred_skills', [])
            top_keywords = jd_skills[:5]  # Top 5 skills
            if top_keywords:
                response += f"**Keywords to include:** {', '.join(top_keywords)}\n\n"

        response += "**Example:** \"Experienced software engineer with 3+ years developing scalable web applications using Python and React. Proven track record of delivering projects on time with 99% uptime. Seeking to leverage expertise in full-stack development to contribute to innovative tech solutions.\"\n\n"

        response += "Would you like me to generate a customized summary for you?"

        return response

    def _handle_experience_request(self, session: Dict[str, Any]) -> str:
        """Handle requests about work experience section"""
        resume = session['resume_json']

        response = "💼 **Work Experience Optimization:**\n\n"

        # Analyze current experience
        experience_years = resume.get('experience_years', 0)
        response += f"You have {experience_years} years of experience.\n\n"

        response += "**Best Practices:**\n"
        response += "• Use strong action verbs (Developed, Implemented, Managed, Created)\n"
        response += "• Quantify achievements (Increased performance by 40%, Managed team of 5)\n"
        response += "• Focus on relevant experience for the target job\n"
        response += "• Keep descriptions concise (2-4 bullet points per role)\n\n"

        response += "**Action Verbs to Use:**\n"
        response += "Developed, Implemented, Created, Optimized, Led, Managed, Designed, Built, Deployed, Improved\n\n"

        if session['jd_json']:
            jd = session['jd_json']
            jd_exp = jd.get('min_experience', 0)
            if experience_years < jd_exp:
                response += f"⚠️ **Gap Alert:** Job requires {jd_exp} years, you have {experience_years}. Consider highlighting relevant projects or skills to compensate.\n\n"

        response += "Would you like me to help rewrite any specific job experience?"

        return response

    def _handle_skills_request(self, session: Dict[str, Any]) -> str:
        """Handle requests about skills section"""
        resume = session['resume_json']
        jd = session['jd_json']

        response = "🛠️ **Skills Section Optimization:**\n\n"

        current_skills = resume.get('tech_skills', [])
        response += f"You currently list {len(current_skills)} technical skills.\n\n"

        if jd:
            jd_mandatory = set(jd.get('mandatory_skills', []))
            jd_preferred = set(jd.get('preferred_skills', []))
            resume_skills = set(current_skills)

            matched_mandatory = resume_skills & jd_mandatory
            missing_mandatory = jd_mandatory - resume_skills
            matched_preferred = resume_skills & jd_preferred

            response += f"**Job Match Analysis:**\n"
            response += f"✅ Matched mandatory skills: {len(matched_mandatory)}\n"
            response += f"❌ Missing mandatory skills: {len(missing_mandatory)}\n"
            response += f"✅ Matched preferred skills: {len(matched_preferred)}\n\n"

            if missing_mandatory:
                response += f"**Skills to Add:** {', '.join(list(missing_mandatory)[:5])}\n\n"

        response += "**Organization Tips:**\n"
        response += "• Group skills by category (Programming Languages, Frameworks, Tools)\n"
        response += "• Put most relevant skills first\n"
        response += "• Include proficiency levels if applicable\n"
        response += "• Keep it to 10-15 most relevant skills\n\n"

        response += "Would you like me to reorganize your skills section?"

        return response

    def _handle_keyword_optimization(self, session: Dict[str, Any]) -> str:
        """Handle ATS keyword optimization requests"""
        jd = session['jd_json']

        response = "🎯 **ATS Keyword Optimization:**\n\n"

        if not jd:
            response += "Please provide a job description for keyword analysis.\n\n"
            response += "**General ATS Tips:**\n"
            response += "• Use industry-standard terms\n"
            response += "• Include exact phrases from job postings\n"
            response += "• Avoid fancy fonts or graphics\n"
            response += "• Use standard section headings\n\n"
            return response

        # Extract keywords from job description
        mandatory = jd.get('mandatory_skills', [])
        preferred = jd.get('preferred_skills', [])

        response += "**High-Priority Keywords (Mandatory):**\n"
        for skill in mandatory[:10]:
            response += f"• {skill}\n"
        response += "\n"

        response += "**Secondary Keywords (Preferred):**\n"
        for skill in preferred[:10]:
            response += f"• {skill}\n"
        response += "\n"

        response += "**ATS Optimization Tips:**\n"
        response += "• Include keywords naturally in your experience descriptions\n"
        response += "• Don't keyword stuff - keep it readable\n"
        response += "• Use exact phrases from the job posting\n"
        response += "• Include both acronyms and full terms (e.g., 'SEO' and 'Search Engine Optimization')\n\n"

        response += "Would you like me to scan your resume for missing keywords?"

        return response

    def _handle_rewrite_request(self, session: Dict[str, Any], message: str) -> str:
        """Handle requests to rewrite specific sections"""
        if 'bullet' in message or 'point' in message:
            bullet = self._extract_first_bullet(session['resume_json'].get('raw_text', ''))
            if bullet:
                return self._rewrite_bullet(bullet)

            return (
                "✍️ **Bullet Rewrite:**\n\n"
                "I can rewrite a bullet point for you, but I need the exact bullet text. "
                "Please paste the line you want me to rewrite."
            )

        response = "✍️ **Content Rewriting:**\n\n"

        # Determine what to rewrite based on message
        if 'summary' in message:
            response += "I'll help you rewrite your professional summary. "
            response += "Please share your current summary, and I'll provide an optimized version."
        elif 'experience' in message:
            response += "I'll help you rewrite your work experience descriptions. "
            response += "Please share a specific job description, and I'll suggest improvements."
        else:
            response += "I can help rewrite any section of your resume. "
            response += "Please specify which section you'd like me to help with (summary, experience, skills, etc.)."

        return response

    def _handle_analysis_request(self, session: Dict[str, Any]) -> str:
        """Provide overall resume analysis"""
        resume = session['resume_json']
        jd = session['jd_json']

        response = "📊 **Resume Analysis Report:**\n\n"

        if jd:
            # Full evaluation
            evaluation = evaluate_resume_against_jd(resume, jd)

            response += f"**Overall Score:** {evaluation['final_score']:.1f}/100\n"
            response += f"**Eligibility:** {evaluation['eligibility']}\n\n"

            # Skills analysis
            skills = evaluation['details']['skills']['mandatory']
            response += f"**Skills Match:** {len(skills['matched'])}/{len(skills['matched']) + len(skills['missing'])} mandatory skills\n"

            # Experience analysis
            exp = evaluation['details']['experience']
            response += f"**Experience:** {exp['found_years']} years (required: {exp['required_years']} years)\n"

            # Education
            edu = evaluation['details']['education']
            response += f"**Education:** {'✅' if edu['pass'] else '❌'} {len(edu['matched'])} requirements met\n"

        else:
            # Basic analysis without JD
            response += "**Basic Resume Health Check:**\n"
            response += f"• Experience: {resume.get('experience_years', 0)} years\n"
            response += f"• Technical Skills: {len(resume.get('tech_skills', []))}\n"
            response += f"• Education: {len(resume.get('education', []))} degrees\n"
            response += f"• Projects: {len(resume.get('projects', []))} detected\n\n"

            response += "**Recommendations:**\n"
            response += "• Add a professional summary if missing\n"
            response += "• Quantify achievements in experience section\n"
            response += "• Include relevant keywords for your target industry\n"

        return response

    def _handle_general_request(self, session: Dict[str, Any], message: str) -> str:
        """Handle general questions and conversation"""
        if any(word in message for word in ['thank', 'thanks']):
            return "You're welcome. I’m here to help with resume improvement, skill gaps, and job readiness. What would you like to refine next?"

        if any(word in message for word in ['help', 'what can you do', 'how']):
            return self._get_help_message()

        evaluation = self._get_evaluation(session)
        if evaluation:
            return (
                f"This resume currently scores {evaluation['final_score']:.1f}/100 against the target role.\n\n"
                "Best next questions:\n"
                "- what are the weaknesses\n"
                "- what skills should I learn\n"
                "- how can I improve my resume\n"
                "- is my resume good for this role"
            )

        return (
            "I can help with resume improvement, skill gaps, ATS keywords, and job readiness.\n\n"
            "Try asking:\n"
            "- how can I improve my resume\n"
            "- what skills should I learn\n"
            "- is my resume good\n"
            "- what are the weaknesses"
        )

    def _get_help_message(self) -> str:
        """Return help message with available features"""
        return """🤖 **Resume Optimizer AI Agent - How I Can Help:**

**📝 Resume Sections:**
• Professional Summary - Make it compelling and keyword-rich
• Work Experience - Use action verbs and quantify achievements
• Skills Section - Organize and prioritize relevant skills
• Education - Highlight relevant qualifications

**🎯 Job-Specific Optimization:**
• ATS Keyword Analysis - Ensure your resume passes applicant tracking systems
• Job Matching - Compare your resume against specific job requirements
• Gap Analysis - Identify missing skills or experience

**✨ Special Features:**
• Content Rewriting - Get AI-suggested improvements
• Section Generation - Create new resume sections from scratch
• Real-time Feedback - Get instant analysis and suggestions

**💬 Just Ask:**
• "Help me improve my summary"
• "Analyze my resume for this job"
• "What keywords should I include?"
• "Rewrite this bullet point"

What would you like to work on first?"""

    def get_suggestions(self, session_id: str) -> List[Dict[str, Any]]:
        """Get optimization suggestions for the current session"""
        if session_id not in self.sessions:
            return []

        session = self.sessions[session_id]
        resume = session['resume_json']
        jd = session['jd_json']

        suggestions = []

        # Basic suggestions
        if not resume.get('summary'):
            suggestions.append({
                'type': 'missing_section',
                'section': 'professional_summary',
                'priority': 'high',
                'message': 'Add a professional summary to introduce yourself to employers',
                'impact': 'High - ATS systems and recruiters look for this first'
            })

        # Job-specific suggestions
        if jd:
            jd_skills = set(jd.get('mandatory_skills', []))
            resume_skills = set(resume.get('tech_skills', []))

            missing_skills = jd_skills - resume_skills
            if missing_skills:
                suggestions.append({
                    'type': 'missing_skills',
                    'section': 'skills',
                    'priority': 'high',
                    'message': f'Consider adding these job-relevant skills: {", ".join(list(missing_skills)[:3])}',
                    'impact': 'High - Direct job requirement match'
                })

        return suggestions

    def generate_section(self, session_id: str, section_type: str, current_content: str = "") -> str:
        """Generate optimized content for a specific resume section"""
        if session_id not in self.sessions:
            return "Session expired. Please start a new session."

        session = self.sessions[session_id]
        jd = session['jd_json']

        if section_type == 'summary':
            return self._generate_summary(session, current_content)
        elif section_type == 'experience':
            return self._generate_experience_description(current_content)
        elif section_type == 'skills':
            return self._generate_skills_section(session)
        else:
            return f"I can help generate content for summary, experience, and skills sections. Please specify one of these."

    def _generate_summary(self, session: Dict[str, Any], current_content: str) -> str:
        """Generate an optimized professional summary"""
        resume = session['resume_json']
        jd = session['jd_json']

        # Extract key information
        experience_years = resume.get('experience_years', 0)
        top_skills = resume.get('tech_skills', [])[:3]
        role = "software engineer"  # This could be detected from resume

        summary = f"Experienced {role} with {experience_years} years of expertise in "

        if top_skills:
            summary += f"{', '.join(top_skills)}. "

        summary += "Proven track record of delivering high-quality solutions and driving project success. "

        if jd:
            jd_title = jd.get('title', 'professional')
            summary += f"Seeking to leverage technical skills to contribute to {jd_title} opportunities."

        return summary

    def _generate_experience_description(self, current_content: str) -> str:
        """Generate optimized experience bullet points"""
        if not current_content:
            return "Please provide your current experience description to optimize."

        # Basic optimization suggestions
        suggestions = [
            "• Start with strong action verbs (Developed, Implemented, Created, Managed)",
            "• Include quantifiable results (increased by 40%, managed team of 5, reduced costs by 30%)",
            "• Focus on achievements rather than just responsibilities",
            "• Keep each bullet point to 1-2 lines",
            "",
            "**Example Optimization:**",
            f"Original: {current_content[:50]}...",
            "Optimized: Developed full-stack web application using React and Node.js, resulting in 50% faster load times and improved user engagement by 35%"
        ]

        return "\n".join(suggestions)

    def _extract_first_bullet(self, text: str) -> Optional[str]:
        """Extract the first bullet-like line from resume raw text."""
        if not text:
            return None

        patterns = [
            r'(?:^|\n)\s*[-*•]\s*(.+)',
            r'(?:^|\n)\s*\d+\.\s*(.+)',
        ]
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1).strip()

        # Fallback: use first sentence with a strong action verb
        for line in text.splitlines():
            cleaned = line.strip()
            if len(cleaned) > 30 and any(word in cleaned.lower() for word in ['developed', 'managed', 'implemented', 'led', 'created', 'designed', 'improved']):
                return cleaned

        return None

    def _rewrite_bullet(self, bullet_text: str) -> str:
        """Rewrite a single bullet point using the resume content."""
        bullet_text = bullet_text.strip()
        optimized = bullet_text

        if re.match(r'^(responsible for|responsibility for|responsible to)', bullet_text, re.I):
            optimized = re.sub(r'^(responsible for|responsibility for|responsible to)\s+', 'Led ', bullet_text, flags=re.I)
        elif not bullet_text[0].isupper():
            optimized = bullet_text[0].upper() + bullet_text[1:]

        optimized = re.sub(
            r'^(Worked on|Worked with|Worked|Contributed to|Helped|Assisted)',
            'Implemented',
            optimized,
            flags=re.I
        )

        return (
            "✍️ **Bullet Rewrite:**\n\n"
            f"Original: {bullet_text}\n\n"
            f"Optimized: {optimized}"
        )

    def _generate_skills_section(self, session: Dict[str, Any]) -> str:
        """Generate organized skills section"""
        resume = session['resume_json']
        jd = session['jd_json']

        skills = resume.get('tech_skills', [])

        if jd:
            # Prioritize job-relevant skills
            jd_skills = set(jd.get('mandatory_skills', []) + jd.get('preferred_skills', []))
            resume_skills = set(skills)

            prioritized = []
            # Job-relevant skills first
            for skill in jd_skills:
                if skill in resume_skills:
                    prioritized.append(skill)

            # Then other skills
            for skill in skills:
                if skill not in jd_skills:
                    prioritized.append(skill)

            skills = prioritized

        # Organize by category
        programming = [s for s in skills if s in ['python', 'java', 'javascript', 'c++', 'c#', 'php', 'ruby', 'go']]
        frameworks = [s for s in skills if s in ['react', 'angular', 'vue', 'django', 'flask', 'spring', 'laravel']]
        tools = [s for s in skills if s in ['git', 'docker', 'kubernetes', 'aws', 'azure', 'linux']]

        organized = ""
        if programming:
            organized += f"**Programming Languages:** {', '.join(programming)}\n"
        if frameworks:
            organized += f"**Frameworks & Libraries:** {', '.join(frameworks)}\n"
        if tools:
            organized += f"**Tools & Technologies:** {', '.join(tools)}\n"

        return organized or ", ".join(skills)

    def cleanup_expired_sessions(self):
        """Remove expired sessions to free memory"""
        current_time = datetime.now()
        expired_sessions = [
            sid for sid, session in self.sessions.items()
            if current_time - session['last_activity'] > self.session_timeout
        ]

        for sid in expired_sessions:
            del self.sessions[sid]

    def build_updated_resume(self, session_id: str, user_instruction: str = "") -> Dict[str, Any]:
        """Create updated resume content from session data and a user instruction."""
        if session_id not in self.sessions:
            raise ValueError("Session expired. Please start a new optimization session.")

        session = self.sessions[session_id]
        resume = session["resume_json"]
        jd = session.get("jd_json")
        template_id = None
        template_profile = self._template_profile(template_id)

        raw_text = resume.get("raw_text", "")
        lines = [line.strip() for line in raw_text.splitlines() if line.strip()]

        def section_lines(section_name: str) -> List[str]:
            pattern = re.compile(rf"^\s*{re.escape(section_name)}\s*$", re.IGNORECASE)
            section_aliases = {
                "education": ["education"],
                "skills summary": ["skills summary", "skills"],
                "experience": ["experience"],
                "projects": ["projects"],
                "honors and awards": ["honors and awards", "honors", "awards"],
                "hackathons & competitions": ["hackathons & competitions", "hackathons", "competitions"],
            }

            aliases = section_aliases.get(section_name.lower(), [section_name.lower()])
            start_index = None
            end_index = None
            for idx, line in enumerate(lines):
                cleaned = re.sub(r"[:\-]\s*$", "", line).strip().lower()
                if any(cleaned == alias for alias in aliases):
                    start_index = idx + 1
                    break
            if start_index is None:
                return []

            heading_aliases = {
                alias for alias_list in section_aliases.values() for alias in alias_list
            }
            for idx in range(start_index, len(lines)):
                cleaned = re.sub(r"[:\-]\s*$", "", lines[idx]).strip().lower()
                if cleaned in heading_aliases and idx > start_index:
                    end_index = idx
                    break

            chunk = lines[start_index:end_index] if end_index else lines[start_index:]
            return [line for line in chunk if line]

        def pick_section_or_fallback(section_name: str, fallback: List[str]) -> List[str]:
            extracted = section_lines(section_name)
            return extracted if extracted else fallback

        def extract_contact_lines() -> Dict[str, str]:
            contact = {
                "portfolio": "",
                "github": "",
                "email": "",
                "mobile": "",
            }
            for line in lines[:10]:
                lowered = line.lower()
                if not contact["email"] and ("@" in line or "email" in lowered):
                    contact["email"] = line.strip()
                elif not contact["mobile"] and ("mobile" in lowered or re.search(r"\+?\d[\d\s\-()]{7,}\d", line)):
                    contact["mobile"] = line.strip()
                elif not contact["github"] and "github" in lowered:
                    contact["github"] = line.strip()
                elif not contact["portfolio"] and ("portfolio" in lowered or "http" in lowered or "www." in lowered):
                    contact["portfolio"] = line.strip()
            return contact

        contact_info = extract_contact_lines()

        name = lines[0] if lines else "Updated Resume"
        contact_lines = lines[1:3] if len(lines) > 1 else []
        experience_years = resume.get("experience_years", 0)
        tech_skills = list(dict.fromkeys(resume.get("tech_skills", [])))
        soft_skills = list(dict.fromkeys(resume.get("soft_skills", [])))
        education = resume.get("education", [])
        projects_count = resume.get("projects", 0)

        highlighted_keywords: List[str] = []
        missing_keywords: List[str] = []
        job_title = "target role"

        if jd:
            highlighted_keywords = list(dict.fromkeys(
                jd.get("mandatory_skills", [])[:6] + jd.get("preferred_skills", [])[:4]
            ))
            job_title = jd.get("job_title") or jd.get("title") or job_title
            missing_keywords = [skill for skill in highlighted_keywords if skill not in tech_skills][:5]

        if user_instruction:
            instruction_text = user_instruction.lower()
            instruction_keywords = [
                keyword for keyword in tech_skills + highlighted_keywords
                if keyword and keyword.lower() in instruction_text
            ]
            for keyword in reversed(instruction_keywords):
                if keyword in highlighted_keywords:
                    highlighted_keywords.remove(keyword)
                highlighted_keywords.insert(0, keyword)

        top_skills = tech_skills[:4] if tech_skills else highlighted_keywords[:4]
        skill_phrase = ", ".join(top_skills) if top_skills else "modern software development"

        summary_lines = section_lines("summary")
        if summary_lines:
            summary = " ".join(summary_lines)
        else:
            summary = (
                f"Results-driven professional with {experience_years} years of experience in {skill_phrase}. "
                f"Known for building reliable solutions, collaborating effectively, and adapting quickly to business needs."
            )
            if jd:
                summary += f" Tailored for {job_title} opportunities with emphasis on {', '.join(highlighted_keywords[:3]) or 'relevant domain skills'}."
        if template_profile["summary_style"] == "compact_header":
            summary = " ".join(summary.split(". ")[:1]).strip()
        elif template_profile["summary_style"] == "concise":
            summary = " ".join(summary.split(". ")[:2]).strip()
        elif template_profile["summary_style"] == "impactful":
            summary = f"{summary} Strong emphasis on measurable impact and leadership."

        skills_lines = pick_section_or_fallback("skills summary", tech_skills[:])
        skills_section = skills_lines[:]
        for keyword in highlighted_keywords:
            if keyword not in skills_section:
                skills_section.append(keyword)

        experience_lines = pick_section_or_fallback("experience", [
            f"Delivered software solutions using {skill_phrase}, with focus on maintainability and business impact.",
            "Improved team delivery through clear communication, problem solving, and ownership of key tasks.",
        ])
        experience_bullets = experience_lines[:]
        if jd and highlighted_keywords and not experience_lines:
            experience_bullets.append(
                f"Aligned project work to target role requirements by emphasizing {', '.join(highlighted_keywords[:3])}."
            )

        projects_section = pick_section_or_fallback("projects", [])
        if not projects_section and projects_count:
            projects_section.append(
                f"Highlighted {projects_count} relevant project(s) to demonstrate hands-on delivery and practical problem solving."
            )

        notes = []

        evaluation = self._get_evaluation(session)
        match_summary = None
        if evaluation:
            skills_detail = evaluation['details']['skills']['mandatory']
            match_summary = {
                "score": evaluation['final_score'],
                "eligibility": evaluation['eligibility'],
                "matched_skills": skills_detail['matched'][:6],
                "missing_skills": skills_detail['missing'][:6],
                "experience_gap": evaluation['details']['experience']['gap_years'],
            }

        return {
            "name": name,
            "raw_lines": lines,
            "contact_info": contact_info,
            "education_lines": section_lines("education"),
            "honors_lines": section_lines("honors and awards"),
            "hackathons_lines": section_lines("hackathons & competitions"),
            "template_id": template_id,
            "template_name": template_profile["name"],
            "section_order": template_profile["section_order"],
            "contact_lines": contact_lines,
            "summary_lines": summary_lines,
            "skills_lines": skills_lines,
            "experience_lines": experience_lines,
            "projects_lines": projects_section,
            "summary": summary,
            "skills": skills_section,
            "soft_skills": soft_skills,
            "education": education,
            "experience_bullets": experience_bullets,
            "projects": projects_section,
            "notes": notes,
            "match_summary": match_summary,
        }

    def export_updated_resume_docx(self, session_id: str, user_instruction: str = "") -> BytesIO:
        """Generate a downloadable DOCX resume for the current optimization session."""
        updated = self.build_updated_resume(session_id, user_instruction)
        section_order = updated.get("section_order", [])

        document = Document()
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(11)

        title = document.add_paragraph()
        title_run = title.add_run(updated["name"])
        title_run.bold = True
        title_run.font.size = Pt(18)
        if updated.get("template_name"):
            title.add_run(f" | {updated['template_name']}").italic = True

        for contact_line in updated["contact_lines"]:
            document.add_paragraph(contact_line)

        def add_heading(text: str):
            document.add_paragraph().add_run(text).bold = True

        summary_lines = updated.get("summary_lines") or [updated["summary"]]
        skills_lines = updated.get("skills_lines") or [", ".join(updated["skills"]) if updated["skills"] else "Add role-relevant technical skills."]
        experience_lines = updated.get("experience_lines") or updated["experience_bullets"]
        projects_lines = updated.get("projects_lines") or updated["projects"]
        section_map = {
            "summary": ("Professional Summary", summary_lines),
            "skills": ("Core Skills", skills_lines),
            "experience": ("Experience Highlights", experience_lines),
            "projects": ("Projects", projects_lines),
            "education": ("Education", [str(item) for item in updated["education"]]),
        }

        used_sections = set()
        for section_key in section_order:
            if section_key not in section_map:
                continue
            heading, items = section_map[section_key]
            used_sections.add(section_key)
            add_heading(heading)
            if section_key == "experience" or section_key == "projects" or section_key == "education":
                for item in items:
                    document.add_paragraph(str(item), style="List Bullet")
            else:
                for item in items:
                    document.add_paragraph(item)

        if updated["soft_skills"] and "skills" in used_sections:
            add_heading("Soft Skills")
            document.add_paragraph(", ".join(updated["soft_skills"]))

        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        return file_stream

    def export_updated_resume_pdf(self, session_id: str, user_instruction: str = "") -> BytesIO:
        """Generate a downloadable PDF resume for the current optimization session."""
        updated = self.build_updated_resume(session_id, user_instruction)
        section_order = updated.get("section_order", [])
        if False:
            return self._export_resume_info_style_pdf_precise(updated)

        doc = fitz.open()
        page = doc.new_page()
        y = 50
        left = 50
        line_height = 18

        def write_line(text: str, font_size: int = 11, bold: bool = False):
            nonlocal y, page
            if y > 760:
                page = doc.new_page()
                y = 50
            font_name = "hebo" if bold else "helv"
            page.insert_text((left, y), text, fontsize=font_size, fontname=font_name)
            y += line_height + (font_size - 11)

        write_line(updated["name"], 18, True)
        for line in updated["contact_lines"]:
            write_line(line)

        summary_lines = updated.get("summary_lines") or [updated["summary"]]
        skills_lines = updated.get("skills_lines") or [", ".join(updated["skills"]) if updated["skills"] else "Add role-relevant technical skills."]
        experience_lines = updated.get("experience_lines") or updated["experience_bullets"]
        projects_lines = updated.get("projects_lines") or updated["projects"]
        section_map = {
            "summary": ("Professional Summary", summary_lines),
            "skills": ("Core Skills", skills_lines),
            "experience": ("Experience Highlights", experience_lines),
            "projects": ("Projects", projects_lines),
            "education": ("Education", [str(item) for item in updated["education"]]),
        }

        for section_key in section_order:
            if section_key not in section_map:
                continue
            heading, lines = section_map[section_key]
            if not lines:
                continue
            y += 8
            write_line(heading, 13, True)
            for line in lines:
                for wrapped in self._wrap_text_for_pdf(line):
                    write_line(f"- {wrapped}" if section_key in {"experience", "projects", "education"} else wrapped)

        if updated["soft_skills"]:
            y += 8
            write_line("Soft Skills", 13, True)
            write_line(", ".join(updated["soft_skills"]))

        file_stream = BytesIO()
        file_stream.write(doc.tobytes())
        file_stream.seek(0)
        doc.close()
        return file_stream

    def _export_resume_info_style_pdf(self, updated: Dict[str, Any]) -> BytesIO:
        """Render a PDF that matches the compact reference-style layout."""
        doc = fitz.open()
        page = doc.new_page(width=595.28, height=841.89)

        left_x = 34
        right_x = 392
        top_y = 50
        page_width = 560
        raw_lines = updated.get("raw_lines", [])

        def draw_rule(y_pos: float):
            page.draw_line((left_x, y_pos), (page_width, y_pos), color=(0, 0, 0), width=0.6)

        def write_text(x: float, y_pos: float, text: str, font_size: int = 11, bold: bool = False, italic: bool = False):
            font_name = "Times-BoldItalic" if bold and italic else "Times-Bold" if bold else "Times-Italic" if italic else "Times-Roman"
            page.insert_text((x, y_pos), text, fontsize=font_size, fontname=font_name, color=(0, 0, 0))
            return y_pos + font_size + 2

        def write_wrapped(x: float, y_pos: float, text: str, font_size: int = 11, bold: bool = False, italic: bool = False, max_chars: int = 88):
            for line in self._wrap_text_for_pdf(text, max_chars=max_chars):
                y_pos = write_text(x, y_pos, line, font_size=font_size, bold=bold, italic=italic)
            return y_pos

        def find_line(keyword: str, fallback: str = "") -> str:
            for line in raw_lines:
                if keyword.lower() in line.lower():
                    return line.strip()
            return fallback

        def lines_after_heading(heading: str, stop_headings: List[str]) -> List[str]:
            start = None
            lower_lines = [line.strip().lower().rstrip(':') for line in raw_lines]
            for idx, line in enumerate(lower_lines):
                if line == heading.lower():
                    start = idx + 1
                    break
            if start is None:
                return []
            stop_set = {item.lower() for item in stop_headings}
            collected = []
            for line in raw_lines[start:]:
                clean = line.strip().lower().rstrip(':')
                if clean in stop_set:
                    break
                if line.strip():
                    collected.append(line.strip())
            return collected

        def render_labeled_lines(items: List[str], start_y: float, label_x: float = 34, content_x: float = 160) -> float:
            y_pos = start_y
            for item in items:
                if ':' in item:
                    label, value = item.split(':', 1)
                    page.insert_text((label_x, y_pos), f"• {label.strip()}:", fontsize=11, fontname="Times-Bold")
                    page.insert_text((content_x, y_pos), value.strip(), fontsize=11, fontname="Times-Roman")
                else:
                    y_pos = write_wrapped(label_x + 10, y_pos, f"• {item}", font_size=11, max_chars=92)
                    continue
                y_pos += 18
            return y_pos

        name_y = write_text(left_x, top_y, updated["name"], font_size=24, bold=True)
        portfolio_line = updated.get("contact_info", {}).get("portfolio", "")
        github_line = updated.get("contact_info", {}).get("github", "")
        email_line = updated.get("contact_info", {}).get("email", "")
        mobile_line = updated.get("contact_info", {}).get("mobile", "")

        if portfolio_line:
            write_text(left_x, name_y - 4, portfolio_line, font_size=13)
        if github_line:
            write_text(left_x, name_y + 12, github_line, font_size=12)

        if email_line:
            write_text(right_x, top_y, email_line, font_size=12)
        if mobile_line:
            write_text(right_x, top_y + 16, mobile_line, font_size=12)

        y = 88

        def section_title(title: str):
            nonlocal y
            y = write_text(left_x, y, title.upper(), font_size=15)
            draw_rule(y - 2)
            y += 8

        section_title("Education")
        education_lines = updated.get("education_lines") or lines_after_heading("education", ["skills summary", "experience", "projects", "honors and awards", "hackathons & competitions"])
        if education_lines:
            first = education_lines[0]
            parts = [p.strip() for p in re.split(r"\s+\|\s+", first) if p.strip()]
            if parts:
                y = write_text(left_x + 10, y, parts[0], font_size=12, bold=True)
                if len(parts) > 1:
                    page.insert_text((right_x + 10, y - 13), parts[1], fontsize=11, fontname="Times-Roman")
                if len(parts) > 2:
                    y = write_wrapped(left_x + 10, y - 2, parts[2], font_size=11, italic=True, max_chars=86)
                if len(parts) > 3:
                    page.insert_text((right_x + 5, y - 13), parts[3], fontsize=11, fontname="Times-Italic")
            for item in education_lines[1:]:
                y = write_wrapped(left_x + 10, y, item, font_size=10, italic=True, max_chars=96)

        y += 8
        section_title("Skills Summary")
        skills_lines = updated.get("skills_lines") or lines_after_heading("skills summary", ["experience", "projects", "honors and awards", "hackathons & competitions"])
        y = render_labeled_lines(skills_lines, y, label_x=34, content_x=160)

        y += 6
        section_title("Experience")
        experience_lines = lines_after_heading("experience", ["projects", "honors and awards", "hackathons & competitions"])
        if not experience_lines:
            experience_lines = updated.get("experience_lines") or [
                "Backend Development Trainee | Myanotomy | Student Developer (Full-time) | July 2025 - Oct 2025",
                "Backend Performance Optimization: Refactored core backend modules to improve performance and efficiently handle large-scale database operations.",
                "REST API Development: Converted legacy components into structured API-first architecture for better scalability and modularity.",
                "Real-Time WebSocket Integration: Implemented Ratchet PHP WebSocket to enable low-latency real-time updates for posts and discussion threads.",
            ]
        if experience_lines:
            header = experience_lines[0]
            header_parts = [p.strip() for p in re.split(r"\s+\|\s+", header) if p.strip()]
            if header_parts:
                page.insert_text((left_x + 10, y), header_parts[0], fontsize=12, fontname="Times-Bold")
                if len(header_parts) > 1:
                    page.insert_text((right_x + 60, y), header_parts[1], fontsize=12, fontname="Times-Roman")
                y += 16
                if len(header_parts) > 2:
                    page.insert_text((left_x + 10, y), header_parts[2], fontsize=11, fontname="Times-Italic")
                if len(header_parts) > 3:
                    page.insert_text((right_x + 35, y), header_parts[3], fontsize=11, fontname="Times-Italic")
                y += 18
            for bullet in experience_lines[1:4]:
                y = write_wrapped(left_x + 18, y, f"• {bullet}", font_size=11, max_chars=88)
        y += 8

        section_title("Projects")
        project_lines = lines_after_heading("projects", ["honors and awards", "hackathons & competitions"])
        if not project_lines:
            project_lines = updated.get("projects_lines") or [
                "AI Resume Builder & ATS Analyzer: Built AI-powered resume intelligence system with resume parsing, job-description similarity scoring, ATS simulation and optimization suggestions. Designed modular backend architecture for scalable deployment.",
                "ChatPayKit – Razorpay + WhatsApp Payment Automation: Designed backend workflow integrating Razorpay Webhooks with Spring Boot and MySQL. Implemented secure signature verification and automated WhatsApp notifications via WhatsApp Cloud API. Ensured transactional database consistency.",
                "CognitoRAG: This project demonstrates how modern AI techniques such as Retrieval-Augmented Generation, semantic search, and local LLMs can be applied to solve real-world information retrieval problems.",
                "3D Interactive Portfolio: Developed a performance-optimized 3D portfolio using React, Three.js, and framer-motion with mobile responsiveness and low-end device rendering optimization.",
            ]
        for project in project_lines[:4]:
            y = write_wrapped(left_x + 10, y, f"• {project}", font_size=11, max_chars=92)
            y += 4

        honors_lines = lines_after_heading("honors and awards", ["hackathons & competitions"])
        if honors_lines:
            y += 6
            section_title("Honors and Awards")
            for honor in honors_lines[:4]:
                y = write_wrapped(left_x + 10, y, f"• {honor}", font_size=11, max_chars=92)

        hack_lines = lines_after_heading("hackathons & competitions", [])
        if hack_lines:
            y += 6
            section_title("Hackathons & Competitions")
            for item in hack_lines[:4]:
                y = write_wrapped(left_x + 10, y, f"• {item}", font_size=11, max_chars=92)

        file_stream = BytesIO()
        file_stream.write(doc.tobytes())
        file_stream.seek(0)
        doc.close()
        return file_stream

    def _export_resume_info_style_pdf_precise(self, updated: Dict[str, Any]) -> BytesIO:
        """Render a compact, reference-like resume PDF with template-safe content only."""
        doc = fitz.open()
        page = doc.new_page(width=595.28, height=841.89)

        raw_lines = updated.get("raw_lines", [])
        contact = updated.get("contact_info") or {}

        left_x = 34
        right_x = 392
        top_y = 50
        rule_x2 = 560
        y = 88

        def draw_rule(y_pos: float):
            page.draw_line((left_x, y_pos), (rule_x2, y_pos), color=(0, 0, 0), width=0.6)

        def write(x: float, y_pos: float, text: str, size: int = 11, bold: bool = False, italic: bool = False):
            font = "Times-BoldItalic" if bold and italic else "Times-Bold" if bold else "Times-Italic" if italic else "Times-Roman"
            page.insert_text((x, y_pos), text, fontsize=size, fontname=font, color=(0, 0, 0))
            return y_pos + size + 2

        def wrap_write(x: float, y_pos: float, text: str, size: int = 11, max_chars: int = 90, bold: bool = False, italic: bool = False):
            for line in self._wrap_text_for_pdf(text, max_chars=max_chars):
                y_pos = write(x, y_pos, line, size=size, bold=bold, italic=italic)
            return y_pos

        def section_block(heading: str):
            nonlocal y
            y = write(left_x, y, heading.upper(), size=15)
            draw_rule(y - 2)
            y += 8

        def pick_lines(names: List[str]) -> List[str]:
            wanted = {n.lower().strip() for n in names}
            collected: List[str] = []
            capture = False
            for line in raw_lines:
                clean = line.strip().lower().rstrip(':')
                if clean in wanted:
                    capture = True
                    continue
                if capture and clean in {
                    "education", "skills summary", "skills", "experience", "projects",
                    "honors and awards", "honors", "awards", "hackathons & competitions",
                    "hackathons", "competitions"
                }:
                    break
                if capture and line.strip():
                    collected.append(line.strip())
            return collected

        # Header
        name = updated.get("name", "Updated Resume")
        y_name = write(left_x, top_y, name, size=24, bold=True)
        if contact.get("portfolio"):
            write(left_x, y_name - 4, contact["portfolio"], size=13)
        if contact.get("github"):
            write(left_x, y_name + 12, contact["github"], size=12)
        if contact.get("email"):
            write(right_x, top_y, contact["email"], size=12)
        if contact.get("mobile"):
            write(right_x, top_y + 16, contact["mobile"], size=12)

        # Education
        section_block("Education")
        edu = updated.get("education_lines") or pick_lines(["education"])
        if edu:
            head = edu[0]
            parts = [p.strip() for p in re.split(r"\s+\|\s+", head) if p.strip()]
            if parts:
                y = write(left_x + 10, y, parts[0], size=12, bold=True)
                if len(parts) > 1:
                    page.insert_text((right_x + 10, y - 13), parts[1], fontsize=11, fontname="Times-Roman")
                if len(parts) > 2:
                    y = wrap_write(left_x + 10, y - 2, parts[2], size=11, italic=True, max_chars=88)
                if len(parts) > 3:
                    page.insert_text((right_x + 5, y - 13), parts[3], fontsize=11, fontname="Times-Italic")
            for line in edu[1:]:
                y = wrap_write(left_x + 10, y, line, size=10, italic=True, max_chars=96)
        y += 8

        # Skills summary
        section_block("Skills Summary")
        skills = updated.get("skills_lines") or pick_lines(["skills summary", "skills"])
        for item in skills:
            if ":" in item:
                label, value = item.split(":", 1)
                page.insert_text((left_x, y), f"• {label.strip()}:", fontsize=11, fontname="Times-Bold")
                page.insert_text((left_x + 122, y), value.strip(), fontsize=11, fontname="Times-Roman")
                y += 18
            else:
                y = wrap_write(left_x + 10, y, f"• {item}", size=11, max_chars=92)
        y += 6

        # Experience
        section_block("Experience")
        exp = updated.get("experience_lines") or pick_lines(["experience"])
        if exp:
            header = exp[0]
            header_parts = [p.strip() for p in re.split(r"\s+\|\s+", header) if p.strip()]
            if header_parts:
                page.insert_text((left_x + 10, y), header_parts[0], fontsize=12, fontname="Times-Bold")
                if len(header_parts) > 1:
                    page.insert_text((right_x + 60, y), header_parts[1], fontsize=12, fontname="Times-Roman")
                y += 16
                if len(header_parts) > 2:
                    page.insert_text((left_x + 10, y), header_parts[2], fontsize=11, fontname="Times-Italic")
                if len(header_parts) > 3:
                    page.insert_text((right_x + 35, y), header_parts[3], fontsize=11, fontname="Times-Italic")
                y += 18
            for bullet in exp[1:]:
                y = wrap_write(left_x + 18, y, f"• {bullet}", size=11, max_chars=88)
        y += 8

        # Projects
        section_block("Projects")
        projects = updated.get("projects_lines") or pick_lines(["projects"])
        for project in projects:
            y = wrap_write(left_x + 10, y, f"• {project}", size=11, max_chars=92)
            y += 4

        honors = updated.get("honors_lines") or pick_lines(["honors and awards", "honors", "awards"])
        if honors:
            y += 6
            section_block("Honors and Awards")
            for honor in honors:
                y = wrap_write(left_x + 10, y, f"• {honor}", size=11, max_chars=92)

        hackathons = updated.get("hackathons_lines") or pick_lines(["hackathons & competitions", "hackathons", "competitions"])
        if hackathons:
            y += 6
            section_block("Hackathons & Competitions")
            for item in hackathons:
                y = wrap_write(left_x + 10, y, f"• {item}", size=11, max_chars=92)

        file_stream = BytesIO()
        file_stream.write(doc.tobytes())
        file_stream.seek(0)
        doc.close()
        return file_stream

    def _wrap_text_for_pdf(self, text: str, max_chars: int = 90) -> List[str]:
        words = text.split()
        if not words:
            return [""]
        lines = []
        current = words[0]
        for word in words[1:]:
            candidate = f"{current} {word}"
            if len(candidate) <= max_chars:
                current = candidate
            else:
                lines.append(current)
                current = word
        lines.append(current)
        return lines
